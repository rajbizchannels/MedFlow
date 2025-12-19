#!/usr/bin/env python3
"""
Convert USER_MANUAL.md to Word format with proper formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document with formatting"""

    # Create a new Document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Configure Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    i = 0
    in_code_block = False
    in_list = False
    list_level = 0

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='Normal')
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # Skip empty lines (but add spacing)
        if not line.strip():
            if i > 0:  # Don't add space at the beginning
                doc.add_paragraph()
            i += 1
            continue

        # Handle H1 headings (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.runs[0]
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.bold = True
            i += 1
            continue

        # Handle H2 headings (## )
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.runs[0]
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.bold = True
            i += 1
            continue

        # Handle H3 headings (### )
        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.bold = True
            i += 1
            continue

        # Handle H4 headings (#### )
        if line.startswith('#### '):
            text = line[5:].strip()
            heading = doc.add_heading(text, level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.runs[0]
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.bold = True
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() in ['---', '___', '***']:
            p = doc.add_paragraph('_' * 80)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.color.rgb = RGBColor(192, 192, 192)
            i += 1
            continue

        # Handle bullet lists (-, *, +)
        if re.match(r'^[\s]*[-*+]\s', line):
            indent_level = len(re.match(r'^[\s]*', line).group()) // 2
            text = re.sub(r'^[\s]*[-*+]\s', '', line)
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            apply_inline_formatting(p)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^[\s]*\d+\.\s', line):
            indent_level = len(re.match(r'^[\s]*', line).group()) // 2
            text = re.sub(r'^[\s]*\d+\.\s', '', line)
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            apply_inline_formatting(p)
            i += 1
            continue

        # Handle blockquotes
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
            i += 1
            continue

        # Handle tables
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Start of a table
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            if len(table_lines) >= 2:
                # Parse table
                header = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                rows = []
                for table_line in table_lines[2:]:  # Skip separator line
                    row = [cell.strip() for cell in table_line.split('|') if cell.strip()]
                    if row:
                        rows.append(row)

                # Create table in Word
                if rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    # Add header
                    for idx, cell_text in enumerate(header):
                        cell = table.rows[0].cells[idx]
                        cell.text = cell_text
                        run = cell.paragraphs[0].runs[0]
                        run.bold = True

                    # Add rows
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(header):
                                table.rows[row_idx + 1].cells[col_idx].text = cell_text

                i = j
                continue

        # Regular paragraph
        text = format_inline_markdown(line)
        p = doc.add_paragraph(text)
        apply_inline_formatting(p)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")

def format_inline_markdown(text):
    """Remove markdown inline formatting markers but keep the text"""
    # This is a simple version - the actual formatting will be applied separately
    return text

def apply_inline_formatting(paragraph):
    """Apply inline formatting to paragraph text"""
    text = paragraph.text
    paragraph.clear()

    # Process the text with inline formatting
    i = 0
    current_text = ""

    while i < len(text):
        # Handle bold (**text** or __text__)
        if (i + 1 < len(text) and text[i:i+2] == '**') or \
           (i + 1 < len(text) and text[i:i+2] == '__'):
            if current_text:
                paragraph.add_run(current_text)
                current_text = ""

            delimiter = text[i:i+2]
            end = text.find(delimiter, i + 2)
            if end != -1:
                bold_text = text[i+2:end]
                run = paragraph.add_run(bold_text)
                run.bold = True
                i = end + 2
                continue

        # Handle italic (*text* or _text_)
        if text[i] in ['*', '_'] and (i == 0 or text[i-1] not in ['*', '_']):
            if current_text:
                paragraph.add_run(current_text)
                current_text = ""

            delimiter = text[i]
            end = text.find(delimiter, i + 1)
            if end != -1 and (end + 1 >= len(text) or text[end+1] != delimiter):
                italic_text = text[i+1:end]
                run = paragraph.add_run(italic_text)
                run.italic = True
                i = end + 1
                continue

        # Handle code (`text`)
        if text[i] == '`':
            if current_text:
                paragraph.add_run(current_text)
                current_text = ""

            end = text.find('`', i + 1)
            if end != -1:
                code_text = text[i+1:end]
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(199, 37, 78)
                i = end + 1
                continue

        current_text += text[i]
        i += 1

    if current_text:
        paragraph.add_run(current_text)

if __name__ == '__main__':
    parse_markdown_to_word('USER_MANUAL.md', 'USER_MANUAL.docx')

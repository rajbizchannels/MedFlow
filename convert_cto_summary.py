#!/usr/bin/env python3
"""
Convert CTO_TECHNICAL_SUMMARY.md to Word format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def parse_markdown_to_word(md_file, docx_file):
    """Convert CTO technical summary to professionally formatted Word document"""

    doc = Document()

    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lang = None

    while i < len(lines):
        line = lines[i]

        # Handle code blocks with language specification
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                code_lang = line.strip()[3:].strip() if len(line.strip()) > 3 else ''
            else:
                code_lang = None
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='Normal')
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            p_format.space_before = Pt(2)
            p_format.space_after = Pt(2)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 51, 102)
            # Add light gray background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            run._element.get_or_add_rPr().append(shading_elm)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            if i > 0:
                doc.add_paragraph()
            i += 1
            continue

        # Handle H1 headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.runs[0]
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0, 32, 96)
            run.bold = True
            i += 1
            continue

        # Handle H2 headings
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            run = heading.runs[0]
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.bold = True
            heading.paragraph_format.space_before = Pt(12)
            i += 1
            continue

        # Handle H3 headings
        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.bold = True
            i += 1
            continue

        # Handle H4 headings
        if line.startswith('#### '):
            text = line[5:].strip()
            heading = doc.add_heading(text, level=4)
            run = heading.runs[0]
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.bold = True
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() in ['---', '___', '***']:
            p = doc.add_paragraph('_' * 80)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.color.rgb = RGBColor(192, 192, 192)
            i += 1
            continue

        # Handle bullet lists
        if re.match(r'^[\s]*[-*+✅❌]\s', line):
            indent_level = len(re.match(r'^[\s]*', line).group()) // 2
            text = re.sub(r'^[\s]*[-*+✅❌]\s', '', line)
            if '✅' in line or '❌' in line:
                symbol = '✅' if '✅' in line else '❌'
                text = symbol + ' ' + text
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            apply_inline_formatting(p)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^[\s]*\d+\.\s', line):
            indent_level = len(re.match(r'^[\s]*', line).group()) // 2
            text = re.sub(r'^[\s]*\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            apply_inline_formatting(p)
            i += 1
            continue

        # Handle tables
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            if len(table_lines) >= 2:
                header = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                rows = []
                for table_line in table_lines[2:]:
                    row = [cell.strip() for cell in table_line.split('|') if cell.strip()]
                    if row:
                        rows.append(row)

                if rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    for idx, cell_text in enumerate(header):
                        cell = table.rows[0].cells[idx]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '002060')
                        cell._element.get_or_add_tcPr().append(shading_elm)

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(header):
                                table.rows[row_idx + 1].cells[col_idx].text = cell_text

                    doc.add_paragraph()
                i = j
                continue

        # Handle bold metadata lines
        if line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            parts = line.split(':**')
            if len(parts) == 2:
                key = parts[0].replace('**', '')
                value = parts[1].strip()
                run1 = p.add_run(key + ': ')
                run1.bold = True
                run2 = p.add_run(value)
            else:
                p.add_run(line)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph(line)
        apply_inline_formatting(p)
        i += 1

    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")

def apply_inline_formatting(paragraph):
    """Apply inline formatting to paragraph text"""
    text = paragraph.text
    paragraph.clear()

    i = 0
    current_text = ""

    while i < len(text):
        # Handle bold
        if (i + 1 < len(text) and text[i:i+2] == '**') or \
           (i + 1 < len(text) and text[i:i+2] == '__'):
            if current_text:
                paragraph.add_run(current_text)
                current_text = ""
            delimiter = text[i:i+2]
            end = text.find(delimiter, i + 2)
            if end != -1:
                bold_text = text[i+2:end]
                run = paragraph.add_run(bold_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                i = end + 2
                continue

        # Handle italic
        if text[i] in ['*', '_'] and (i == 0 or text[i-1] not in ['*', '_']):
            if current_text:
                paragraph.add_run(current_text)
                current_text = ""
            delimiter = text[i]
            end = text.find(delimiter, i + 1)
            if end != -1 and (end + 1 >= len(text) or text[end+1] != delimiter):
                italic_text = text[i+1:end]
                run = paragraph.add_run(italic_text)
                run.italic = True
                i = end + 1
                continue

        # Handle code
        if text[i] == '`':
            if current_text:
                paragraph.add_run(current_text)
                current_text = ""
            end = text.find('`', i + 1)
            if end != -1:
                code_text = text[i+1:end]
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(199, 37, 78)
                i = end + 1
                continue

        current_text += text[i]
        i += 1

    if current_text:
        paragraph.add_run(current_text)

if __name__ == '__main__':
    parse_markdown_to_word('CTO_TECHNICAL_SUMMARY.md', 'CTO_TECHNICAL_SUMMARY.docx')
